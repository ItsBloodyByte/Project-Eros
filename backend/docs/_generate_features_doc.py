"""Generates /app/backend/docs/eros_features.docx - a structured overview of all features."""
import os
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = os.path.join(os.path.dirname(__file__), "eros_features.docx")

SECTIONS = [
    ("1. Authentifizierung & Konto", [
        "Registrierung mit E-Mail, Passwort, Anzeigename, Alter (>= 18, nach Registrierung nicht mehr aenderbar) und Geschlecht (Pflicht)",
        "Zweier-Registrierung als Paar-Account (Duo): Persona A + Persona B in einem Durchlauf, eine Account-ID mit gemeinsamer Sichtbarkeit in Entdecken",
        "Partner-Einladung nachtraeglich per Invite-Link (POST /api/auth/invite-partner), getrennte Logins, gekoppelte Profile",
        "Login mit JWT-Token + automatische Wiederherstellung der Sitzung bei Seiten-Refresh",
        "E-Mail-Verifizierung mit Dev-Code (Produktion: Mail-Versand vorgesehen)",
        "Zwei-Faktor-Authentifizierung (TOTP) mit eigener Aktivierung/Deaktivierung",
        "Passwort-Mindestanforderung: 8+ Zeichen",
        "Einwilligungen (Nutzungsbedingungen, Datenschutz, sensible Daten, optional NSFW-Ansicht)",
        "Logout, Theme- und Sprach-Persistenz in LocalStorage",
    ]),
    ("2. Profil & Selbstdarstellung", [
        "Editierbares Profil: Anzeigename, Bio, Pronomen, Orientierung, Beziehungstypen, gesuchte Rollen, Kinks",
        "Erweiterte Attribute: Koerpergroesse, Koerpertyp, Ethnie, Sprachen, Interessen, Rauch-/Trink-/Diaet-Verhalten, STI-Status mit Datum, Stadt (sichtbar in Profilkarten und Detailansicht)",
        "Geschlechts-bedingte Felder: Cup-Size oder Penis-Laenge/-Umfang (mit automatischer S/M/L/XL/XXL-Kategorisierung)",
        "Bis zu 5 Fotos (1 Haupt + 4 Neben), Drag-and-Drop zum Umsortieren auf Desktop, Als-Hauptfoto-Button auf Mobile",
        "Profilfoto-Verwaltung zusaetzlich im Alben-Bereich (Tab 'Profilfotos') mit denselben Controls",
        "Serverseitige Foto-Kompression beim Upload: Pillow-basiert, max. 1600 px Kantenlaenge, JPEG q82, EXIF-strip, RGBA-Flattening (typisch -80 Prozent Wire-Groesse)",
        "NSFW- und Gesichts-Erkennung per Gemini-Vision pro Foto (Score + Flag)",
        "Blur-Overlay bei NSFW-Bildern mit expliziter Freigabe durch den Betrachter",
        "Videos: Upload und Wiedergabe auf dem Profil",
        "Alter ist nach Registrierung immutable (serverseitig silent-ignored)",
        "Persona-B-Editor fuer Duo-Accounts: eigener Anzeigename, Foto, Bio, Gender, Orientierung, Kinks; Persona-B-Daten in Card und Profilansicht sichtbar",
    ]),
    ("3. Status-Indikatoren (Moods)", [
        "Vier Preset-Status: Sex-Treffen (Flame, destruktive Farbe), Dating (Heart, Akzent), Chatten (Message-Circle, Ring-Farbe), Online (Sparkles, Emerald)",
        "Schnelles Toggling via PATCH /api/me/mood mit Zeitstempel-Tracking; Clear via null",
        "MoodSelector auf eigenem Profil: Karten-UI mit Icon, Label, Beschreibung und Status-Entfernen-Button",
        "Mood-Badge auf Desktop-Karte: volle Pille mit Icon + Label an der linken unteren Kante",
        "Mood-Badge auf Mobile-Karte: Icon-Only in der gleichen Groesse wie das Augen-Symbol, platziert im Top-Right-Cluster neben Online-Dot",
        "Mood-Badge zusaetzlich in ProfileView (Header, Groesse md) und im Chat-Peer-Header (Icon-Only auf Mobile, Pille auf Desktop)",
        "Mood-Filter im Filter-Drawer: Chip-Auswahl aller vier Stati, filtert /api/discover per preferences.moods",
    ]),
    ("4. Entdecken (Discover) & Suche", [
        "Poster-artiges Grid mit grossen Profilkarten",
        "Tageszeit-abhaengige Begruessung im Hero: Guten Morgen / Guten Tag / Guten Abend / Gute Nacht (tageszeit-lokal berechnet)",
        "Quick-Filter-Chip-Strip: Online, Verifiziert, Mit-Gesicht, Mit-Foto",
        "Erweiterter Filter-Drawer: Alter, Entfernung, Geschlecht, Status, Beziehungstyp, Groesse, Koerpertyp, Sprachen, Kinks, Gewohnheiten, STI-Status, Cup-/Penis-Kategorien",
        "Inklusive Gender-Match-Logik: Profil wird angezeigt, wenn ENTWEDER der Betrachter das Gender sucht ODER die Zielperson das Gender des Betrachters sucht (kein strenger bidirektionaler AND-Filter mehr)",
        "Bidirektionale Alters-Filterung bleibt erhalten",
        "Bereits-angesehen-Markierung per Augen-Icon; Profile bleiben sichtbar",
        "Boost-Sortierung: aktiv geboostete Profile erscheinen oben",
        "Pagination mit Mehr-Laden-Button, Skeleton-Ladezustand und Empty-State",
        "Distanzberechnung per Geo-2dsphere-Index (haversine)",
        "Admin-Ansicht (Staff-only) mit Bulk-Aktionen: Hide/Unhide, Ban/Unban fuer gewaehlte Profile",
        "Paar-Accounts: beide Personas sichtbar mit 'PAAR'-Badge, Partner-Snapshot angehaengt",
        "List-Mode-Optimierung: List-Endpoints senden nur das Primaerfoto statt aller 5 (drastische Wire-Groessen-Reduktion)",
    ]),
    ("5. Views (Profil-Besucher:innen)", [
        "Eigene Seite /visitors mit Grid-Layout, Skeleton, Empty-/Error-State",
        "Premium-Ansicht: alle Views innerhalb 30 Tage, vollstaendige Karte mit Name, Alter, Zeitstempel, Visit-Count",
        "Free-Tier: letzte 3 Views klar sichtbar, darueber hinausgehende Views innerhalb der letzten 24 Stunden als geblurte Silhouetten-Karten mit Premium-Lock-Badge",
        "Premium-CTA unten mit dynamischer Zaehlung (X verschleierte Views, Upgrade-Link)",
        "Anklickbare View-Karten fuehren direkt zum Profil (zaehlt nicht als weiterer Besuch)",
        "Backend-Endpoint GET /api/me/unread-summary fuer Navigation-Counter (unread_messages, unread_matches, new_matches) im 30s-Polling-Hook",
    ]),
    ("6. Interaktion, Matches & Likes", [
        "Like / Pass (1:1 Like-Tabelle mit Unique-Index)",
        "Tagliches Free-Like-Limit von 5, Premium unlimitiert",
        "Super-Likes mit eigenem Tageskontingent, priorisierte Anzeige",
        "Visitor-Tracking (jeder Profilbesuch wird deduplikiert gezaehlt)",
        "Match-Erkennung bei gegenseitigem Like",
        "Wer-mich-geliked-hat (Premium-only)",
        "Erst-Nachricht ohne Match (Premium-only)",
        "Unmatch / Block-Funktion",
        "Melde-Funktion mit Gruenden; Zieltypen user oder message; Bilder im Report mit Vorschau sichtbar",
        "Report-Lock: waehrend aktiver Meldung blockiert der Server das Loeschen betroffener Fotos (423 Locked) und zeigt klaren Hinweis im UI",
    ]),
    ("7. Chat & Messaging", [
        "Echtzeit-Chat via WebSocket (Typing, Presence, Screenshot-Events)",
        "Chat-Historie mit Lesebestaetigungen (Single-/Double-Check)",
        "Asymmetrische Nachrichtenblasen (eigener Akzent vs. neutrale Karte)",
        "Klickbarer Chat-Peer-Header: Avatar + Name fuehren zur Profilseite (Ausnahme: System-/Offizielle Accounts)",
        "Peer-Header zeigt Paar-Badge, Mood-Icon (responsiv), Online-Dot und Offiziell-Siegel",
        "Duo-Chat: Multi-Sender-UI, Nachrichten beider Personas unterscheidbar",
        "Broadcast-Nachrichten erscheinen als Chat im Postfach (System-/Offizielles Profil)",
        "Medien-Upload im Chat (Bilder) mit NSFW-Blur bei Score >= 0,75",
        "Selbstzerstoerende Nachrichten (60 Sekunden)",
        "Link-Blocker: lehnt URLs, (dot)/[dot]-Obfuskierung, leerzeichen-getrennte Domains (google com) und buchstaben-getrennte Varianten (g o o g l e . c o m) sowie E-Mail-Adressen ab; ohne Falschpositive bei deutschen Texten",
        "Privatsphaere-Toggles im Chat-Menue (Lesebestaetigungen)",
    ]),
    ("8. Broadcasts (segmentiert)", [
        "Admin-Composer mit Segmentierung: Rolle, Premium-Status, Account-Typ, Gender, Stadt, Altersbereich, Status, Kinks, etc.",
        "Live-Count der erreichten Empfaenger beim Komponieren (nicht-persistierte Vorschau)",
        "Versand erzeugt Chat-Nachrichten vom System-Account ins Postfach jedes Empfaengers",
        "Broadcast-Banner global verfuegbar fuer wichtige Ankuendigungen",
    ]),
    ("9. Premium & Zahlungen", [
        "Premium-Status mit Ablaufdatum (verlaengernd)",
        "Boost-Pakete (Minuten-basiert, sortieren Discover nach oben)",
        "Dynamisch konfigurierbare Pakete (ID, Betrag, Waehrung, Art, Dauer, Beschreibung, aktiv/inaktiv)",
        "Multi-Provider-Konfiguration: Stripe (live-integriert), PayPal, Mollie, Klarna, Paddle, Custom",
        "Admin-UI zur Verwaltung der Anbieter-Schluessel (maskierte Rueckgabe) und Pakete",
        "Stripe-Checkout-Session-Erzeugung, Webhook, Status-Endpoint, idempotente Gutschrift",
        "Premium-Badge im Header bei aktivem Abo",
        "Boost-Badge sichtbar solange aktiv",
    ]),
    ("10. Promo-Codes", [
        "Admin-Tab 'Promos' zum Erstellen/Verwalten von Codes",
        "Code-Attribute: max_uses, used_count, valid_until, premium_days (gutgeschriebene Tage)",
        "Einloeseflow im Konto-Bereich: Nutzer gibt Code ein, Server validiert + schreibt Premium-Tage gut",
        "Audit-Eintraege bei Einloesungen fuer Transparenz",
    ]),
    ("11. Blog (Rich-Text)", [
        "Admin-Editor mit TipTap (Starter-Kit): Headings, Listen, Quotes, Bold/Italic, Hyperlinks, Code, Bilder",
        "Beitraege besitzen slug, title, html_content, published_at, tags",
        "Oeffentliche Liste /blog mit Teasern und Tag-Filter",
        "Einzelseite /blog/[slug] mit editorialer Typografie (Playfair + Figtree + Prose)",
        "Admin-Actions: Draft / Publish / Archive / Delete",
    ]),
    ("12. ID-Verifizierung (kostenlos)", [
        "Nutzer-Upload von Selfie + Dokument (Reisepass / Personalausweis / Fuehrerschein)",
        "Status-Flow: pending -> approved / rejected",
        "Admin-Review-Queue mit Thumbnail-Anzeige und Entscheidungs-Aktionen",
        "Verifiziert-Badge (blauer Schild) in Profilkarten und Profilansicht",
    ]),
    ("13. Reisen & Events", [
        "Reiseplanung: Ziel-Stadt/Land, Zeitraum, optionale Geo-Koordinaten",
        "Liste eigener Reisen mit Loesch-Option",
        "Events-System: Auflistung, RSVP, Event-Details",
    ]),
    ("14. Alben (privat / geteilt)", [
        "Tab-basierte Alben-Seite: 'Profilfotos' (Verwaltung der 5 Profilbilder) und 'Private Alben'",
        "Private Alben mit Unlock-Mechanik (Freigabe pro Nutzer)",
        "Album-CRUD, Unlock-Anfragen und Genehmigungen",
    ]),
    ("15. Moderation & Anti-Abuse", [
        "KI-gestuetzte Bildpruefung (Gemini Vision) beim Upload",
        "Admin-Foto-Queue mit Threshold-Filter",
        "Granulare Moderator-Kanaele fuer Teams / Rollenzuweisung",
        "Auto-Moderation: nach 10 eindeutigen Melde-IPs -> Shadow-Restrict des Zielusers",
        "Duplikat-Schutz bei Meldungen (reporter/target)",
        "Manuelle Aktionen: Ban/Unban, Shadow-Restrict, Warnungen",
        "Foto-Retention: Admin kann pro Foto eine 30-Tage-Aufbewahrung markieren (auch wenn Nutzer loescht)",
        "Waehrend aktiver Reports: Nutzer-Foto-Loeschung serverseitig blockiert",
        "Audit-Log fuer alle Admin-Aktionen",
    ]),
    ("16. Admin-Panel (Backoffice)", [
        "Tabs: Reports, Users, Foto-Queue, Verifizierungen, KI-Konfig, Zahlungen, Promos, Blog, Broadcasts, Rechtliches (CMS), Audit",
        "Benutzersuche, Status-Badges (banned / shadow / ID / aktiv)",
        "Bulk-User-Actions: Selektiere mehrere Nutzer und fuehre Hide/Unhide/Ban/Unban aus",
        "Rollenbadges (Team / Staff / Moderator / Admin / Superadmin) mit Hover-Tooltip zur Erklaerung",
        "KI-Konfiguration: Provider (Gemini / OpenAI / Ollama), Modell, API-Key, Aktiv-Flag",
        "Zahlungs-Konfiguration mit Anbieter-Cards und Package-Editor",
        "Rechtliches (CMS light): 6 Seiten editierbar als Markdown mit Live-Vorschau",
        "Rollen: user / moderator / admin / superadmin mit Sichtbarkeits-Gating der Tabs",
    ]),
    ("17. Rechtliche Seiten & Transparenz", [
        "Oeffentlich lesbar ohne Login: /legal/terms, /legal/privacy, /legal/imprint, /legal/community, /legal/cookies, /legal/cancellation",
        "Markdown-Rendering mit editorialer Typografie",
        "Footer mit Legal-Links auf Login, Discover, Account und anderen Seiten",
        "Deutsche Default-Inhalte, editierbar durch Admin",
    ]),
    ("18. Datenschutz & DSGVO", [
        "Daten-Export (JSON-Download aller eigenen Daten)",
        "Konto-Loeschung mit Bestaetigungs-Dialog und vollstaendigem Entfernen",
        "Privatsphaere-Schalter: Lesebestaetigungen, Online-Status, Tipp-Indikator, Versteckter Modus, Screenshot-Hinweise",
        "Screenshot-Deterrent: CSS-Hinweise, Druck-Block, Blur bei erkanntem Screenshot-Event",
        "Keine Tracking-Scripts (Posthog / Emergent-Tracker entfernt)",
    ]),
    ("19. Navigation & Informationsarchitektur", [
        "Desktop-Header: Primaer-Nav (Entdecken, Matches mit Counter, Views, Alben, Events, Blog) + Menue-Dropdown (Konto, Einstellungen, Admin, Sprache-Submenue, Abmelden)",
        "Mobile: Hamburger-Sidemenu mit kompletten Navigationsgruppen + Konto-Gruppe + Einstellungen (Theme, Sprache) + Logout",
        "Mobile-Bottom-Nav (sticky): Entdecken | Views | Matches | Alben (4-Tab), Icons mit Akzent-Farbe bei aktiver Route",
        "Globaler Unread-Counter-Hook (useUnreadCounts), /api/me/unread-summary mit unread_messages + new_matches (letzte 72h); 30s-Polling mit visibilitychange-Trigger",
        "Matches-Badge erscheint in Desktop-Nav und Mobile-Bottom-Nav, verschwindet bei Count = 0",
    ]),
    ("20. Internationalisierung", [
        "Deutsch (primaer) + Englisch via react-i18next",
        "Sprachwechsler im Menue-Dropdown (Desktop) bzw. Sheet (Mobile) mit Persistenz in LocalStorage",
        "Formulare, Filter, Navigation und Fehler-Toasts alle lokalisiert",
    ]),
    ("21. Design-System & UX", [
        "Dark- und Light-Theme gleichwertig, System-Preference-Auto-Detection, manuell umschaltbar",
        "Akzentfarbe Apricot (warm, ruhig) mit Verified-Blau als Sekundaer-Akzent",
        "Typografie: Playfair Display (Editorial) + Figtree (Body) + IBM Plex Mono",
        "Abgerundete Karten (radius-lg), sanfte Schatten, Ring-basierter Rahmen",
        "Sticky Glass-Navigation mit aktiver Pille, Hover/Active-Mikrointeraktionen",
        "Mobile-first-Layout, 44x44 px Touch-Targets, prefers-reduced-motion-Fallback",
        "Skeleton-Loader, Empty-/Error-States, Sonner-Toasts fuer Feedback",
    ]),
    ("22. Performance", [
        "GZip-Middleware (compresslevel 6, min 1 KB) transparent auf allen JSON-Endpoints",
        "Pillow-basierte Upload-Kompression: Resize, JPEG q82, Alpha-Flattening",
        "One-Shot-Migration fuer bestehende Fotos (typisch 80+ Prozent Ersparnis)",
        "list_mode-Parameter in public_user_from_doc: nur Primaerfoto in List-Responses (/discover, /matches, /visitors, Admin-Discover)",
        "Messbare Ergebnisse: /me von 7.1 MB auf 276 KB (-96 Prozent), /discover 2.3 MB auf 321 KB (-86 Prozent), /matches 2 MB auf 132 KB (-93 Prozent)",
        "Initial-Load-Zeiten durchgehend unter 200 ms",
    ]),
    ("23. Sicherheit & Architektur", [
        "FastAPI + MongoDB (motor async) mit UUIDs als IDs",
        "Timezone-aware datetimes (UTC)",
        "JWT mit Secret-Env, Axios-Interceptor bei 401 -> /login",
        "CORS konfigurierbar, WebSocket-Token-Auth",
        "Pydantic-Modelle fuer alle Requests/Responses (inkl. Mood-Literal, MoodUpdateRequest)",
        "Supervisor-managed Services (Frontend / Backend)",
        "Emergent LLM Universal Key fuer Gemini-Vision (keine direkte SDK-Abhaengigkeit)",
    ]),
    ("24. Mobile App (React Native / Expo)", [
        "Iteration 1 abgeschlossen: Auth-Flow, Theme-Switch, Matches-Liste, Chat-Basics",
        "Iteration 2 in Planung: Feature-Parity mit Filtern, Uploads, Blog, Events, Alben, Stealth/Visitors, Video, Couples-UI",
    ]),
]


def add_title(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(28)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(15)
        run.font.color.rgb = RGBColor(0x2a, 0x2a, 0x2a)
    else:
        run.font.size = Pt(11)


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    add_title(doc, "Eros \u2014 Feature- und Funktionsuebersicht", level=0)
    sub = doc.add_paragraph()
    sub_run = sub.add_run(f"Stand: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "Dieses Dokument listet alle aktuell implementierten Features und "
        "Funktionen der Eros-Plattform. Gliederung nach Funktionsbereichen. "
        "ID-Verifizierung ist kostenlos. Zahlungen sind multi-provider-faehig, "
        "wobei aktuell nur Stripe live integriert ist."
    )
    for run in intro.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()

    for heading, items in SECTIONS:
        add_title(doc, heading, level=1)
        for item in items:
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(item)
            run.font.size = Pt(10.5)
        doc.add_paragraph()

    footer = doc.add_paragraph()
    fr = footer.add_run(
        "Dieses Dokument wird beim Ausfuehren von backend/docs/_generate_features_doc.py neu erzeugt."
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
